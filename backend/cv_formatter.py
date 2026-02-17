"""
cv_formatter.py — CV Maltem Africa — 100% fidèle référence ZAID v2
"""
import os, re
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

RED   = RGBColor(0xC0, 0x00, 0x00)
TITLE = RGBColor(0xBE, 0x3B, 0x4E)
BLACK = RGBColor(0x11, 0x11, 0x11)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREY  = RGBColor(0x60, 0x60, 0x60)

ASSETS  = os.path.join(os.path.dirname(__file__), "assets")
LOGO    = os.path.join(ASSETS, "logo_maltem.png")
DECO_TL = os.path.join(ASSETS, "deco_top_left_white.png")
DECO_BR = os.path.join(ASSETS, "deco_bottom_right_white.png")


# ─── UTILITAIRES ──────────────────────────────────────────────────────────────

def sf(run, size_pt=10, bold=False, italic=False, underline=False, color=None, font="Century Gothic"):
    run.font.name = font
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    run.font.underline = underline
    if color:
        run.font.color.rgb = color
    rPr = run._r.get_or_add_rPr()
    rF = OxmlElement('w:rFonts')
    for a in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
        rF.set(qn(a), font)
    ex = rPr.find(qn('w:rFonts'))
    if ex is not None:
        rPr.remove(ex)
    rPr.insert(0, rF)


def sp(para, before=0, after=0, line=240):
    pPr = para._p.get_or_add_pPr()
    s = pPr.find(qn('w:spacing'))
    if s is None:
        s = OxmlElement('w:spacing')
        pPr.append(s)
    s.set(qn('w:before'), str(int(before * 20)))
    s.set(qn('w:after'),  str(int(after  * 20)))
    s.set(qn('w:line'),   str(line))
    s.set(qn('w:lineRule'), 'auto')


def ind(para, left=0, hanging=0):
    pPr = para._p.get_or_add_pPr()
    i = pPr.find(qn('w:ind'))
    if i is None:
        i = OxmlElement('w:ind')
        pPr.append(i)
    if left:
        i.set(qn('w:left'), str(left))
    if hanging:
        i.set(qn('w:hanging'), str(hanging))


def bdr(para, top=False, bottom=False, color="231F20", sz="6"):
    pPr = para._p.get_or_add_pPr()
    ex = pPr.find(qn('w:pBdr'))
    if ex is not None:
        pPr.remove(ex)
    pBdr = OxmlElement('w:pBdr')
    for side, do in [('top', top), ('bottom', bottom)]:
        if do:
            e = OxmlElement(f'w:{side}')
            e.set(qn('w:val'),   'single')
            e.set(qn('w:sz'),    sz)
            e.set(qn('w:space'), '1')
            e.set(qn('w:color'), color)
            pBdr.append(e)
    pPr.append(pBdr)


def page_break_before(para):
    """Forcer un saut de page AVANT ce paragraphe."""
    pPr = para._p.get_or_add_pPr()
    pb = OxmlElement('w:pageBreakBefore')
    pb.set(qn('w:val'), '1')
    pPr.append(pb)


def remove_cell_borders(cell):
    """Supprimer toutes les bordures d'une cellule."""
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        tc.insert(0, tcPr)
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'none')
        b.set(qn('w:sz'), '0')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), 'auto')
        tcBorders.append(b)
    old = tcPr.find(qn('w:tcBorders'))
    if old is not None:
        tcPr.remove(old)
    tcPr.append(tcBorders)


def set_cell_width(cell, width_twips):
    """Définir la largeur d'une cellule en twips."""
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        tc.insert(0, tcPr)
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'),    str(width_twips))
    tcW.set(qn('w:type'), 'dxa')
    old = tcPr.find(qn('w:tcW'))
    if old is not None:
        tcPr.remove(old)
    tcPr.append(tcW)


# ─── IMAGES / FORMES FLOTTANTES ───────────────────────────────────────────────

def make_anchor_from_inline(run, cx, cy, posH_offset, posV_offset, doc_id,
                             relH="column", relV="paragraph", behindDoc="0"):
    r_xml = run._r.xml
    m = re.search(r'r:embed="(rId\d+)"', r_xml)
    if not m:
        return None
    rId = m.group(1)
    anchor_xml = f'''<w:drawing
  xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
  xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
  xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
  xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
  <wp:anchor distT="0" distB="0" distL="114300" distR="114300" simplePos="0"
    relativeHeight="251658240" behindDoc="{behindDoc}" locked="0" layoutInCell="1" allowOverlap="1">
    <wp:simplePos x="0" y="0"/>
    <wp:positionH relativeFrom="{relH}"><wp:posOffset>{posH_offset}</wp:posOffset></wp:positionH>
    <wp:positionV relativeFrom="{relV}"><wp:posOffset>{posV_offset}</wp:posOffset></wp:positionV>
    <wp:extent cx="{cx}" cy="{cy}"/>
    <wp:effectExtent l="0" t="0" r="0" b="0"/>
    <wp:wrapNone/>
    <wp:docPr id="{doc_id}" name="img{doc_id}"/>
    <a:graphic>
      <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">
        <pic:pic xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture">
          <pic:nvPicPr>
            <pic:cNvPr id="0" name="img{doc_id}"/>
            <pic:cNvPicPr><a:picLocks noChangeAspect="1"/></pic:cNvPicPr>
          </pic:nvPicPr>
          <pic:blipFill>
            <a:blip r:embed="{rId}"/>
            <a:stretch><a:fillRect/></a:stretch>
          </pic:blipFill>
          <pic:spPr>
            <a:xfrm><a:off x="0" y="0"/><a:ext cx="{cx}" cy="{cy}"/></a:xfrm>
            <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>
          </pic:spPr>
        </pic:pic>
      </a:graphicData>
    </a:graphic>
  </wp:anchor>
</w:drawing>'''
    return etree.fromstring(anchor_xml)


def make_red_bar_xml():
    return etree.fromstring('''<w:drawing
  xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
  xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
  xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
  xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
  <wp:anchor distT="0" distB="0" distL="114300" distR="114300" simplePos="0"
    relativeHeight="251659264" behindDoc="1" locked="0" layoutInCell="1" allowOverlap="1">
    <wp:simplePos x="0" y="0"/>
    <wp:positionH relativeFrom="page"><wp:align>left</wp:align></wp:positionH>
    <wp:positionV relativeFrom="paragraph"><wp:posOffset>-647700</wp:posOffset></wp:positionV>
    <wp:extent cx="7560310" cy="72390"/>
    <wp:effectExtent l="0" t="0" r="2540" b="3810"/>
    <wp:wrapNone/>
    <wp:docPr id="1" name="RedBar"/>
    <a:graphic>
      <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
        <wps:wsp>
          <wps:cNvSpPr><a:spLocks noChangeArrowheads="1"/></wps:cNvSpPr>
          <wps:spPr bwMode="auto">
            <a:xfrm><a:off x="0" y="0"/><a:ext cx="7560310" cy="72390"/></a:xfrm>
            <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>
            <a:solidFill><a:srgbClr val="E9272D"/></a:solidFill>
            <a:ln><a:noFill/></a:ln>
          </wps:spPr>
          <wps:bodyPr rot="0" wrap="square" anchor="t"><a:noAutofit/></wps:bodyPr>
        </wps:wsp>
      </a:graphicData>
    </a:graphic>
  </wp:anchor>
</w:drawing>''')


def make_badge_xml(periode_text, doc_id=100):
    cx = 1097280
    cy = 246888
    base_xml = f'''<w:drawing
  xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
  xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
  xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
  xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
  <wp:anchor distT="0" distB="0" distL="114300" distR="114300" simplePos="0"
    relativeHeight="251664384" behindDoc="0" locked="0" layoutInCell="1" allowOverlap="1">
    <wp:simplePos x="0" y="0"/>
    <wp:positionH relativeFrom="column"><wp:posOffset>-50800</wp:posOffset></wp:positionH>
    <wp:positionV relativeFrom="paragraph"><wp:posOffset>100000</wp:posOffset></wp:positionV>
    <wp:extent cx="{cx}" cy="{cy}"/>
    <wp:effectExtent l="0" t="0" r="1270" b="5715"/>
    <wp:wrapNone/>
    <wp:docPr id="{doc_id}" name="badge{doc_id}"/>
    <a:graphic>
      <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
        <wps:wsp>
          <wps:cNvSpPr><a:spLocks/></wps:cNvSpPr>
          <wps:spPr>
            <a:xfrm><a:off x="0" y="0"/><a:ext cx="{cx}" cy="{cy}"/></a:xfrm>
            <a:prstGeom prst="roundRect"><a:avLst><a:gd name="adj" fmla="val 16667"/></a:avLst></a:prstGeom>
            <a:solidFill><a:srgbClr val="C00000"/></a:solidFill>
            <a:ln><a:noFill/></a:ln>
          </wps:spPr>
          <wps:txbx>
            <w:txbxContent>
              <w:p>
                <w:pPr><w:jc w:val="center"/>
                  <w:rPr>
                    <w:rFonts w:ascii="Century Gothic" w:hAnsi="Century Gothic" w:cs="Century Gothic"/>
                    <w:color w:val="FFFFFF"/><w:sz w:val="16"/><w:szCs w:val="16"/>
                  </w:rPr>
                </w:pPr>
                <w:r>
                  <w:rPr>
                    <w:rFonts w:ascii="Century Gothic" w:hAnsi="Century Gothic" w:cs="Century Gothic"/>
                    <w:color w:val="FFFFFF"/><w:sz w:val="16"/><w:szCs w:val="16"/>
                  </w:rPr>
                  <w:t xml:space="preserve">BADGE_PLACEHOLDER</w:t>
                </w:r>
              </w:p>
            </w:txbxContent>
          </wps:txbx>
          <wps:bodyPr rot="0" anchor="ctr"><a:noAutofit/></wps:bodyPr>
        </wps:wsp>
      </a:graphicData>
    </a:graphic>
  </wp:anchor>
</w:drawing>'''
    root = etree.fromstring(base_xml)
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    for wt in root.iter(f"{{{W}}}t"):
        if wt.text == "BADGE_PLACEHOLDER":
            wt.text = periode_text
            break
    return root


# ─── EN-TÊTE ──────────────────────────────────────────────────────────────────

def build_header(doc, cv_data):
    p_main = doc.add_paragraph()
    p_main.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sp(p_main, before=0, after=0)

    r_empty = p_main.add_run(" ")
    sf(r_empty, size_pt=10)
    p_main._p.insert(1, make_red_bar_xml())

    r_logo = p_main.add_run()
    r_logo.add_picture(LOGO, width=Cm(6.76))
    logo_anchor = make_anchor_from_inline(r_logo, 2433960, 635257, 4852434, -482157, 44,
                                          relH="page", relV="paragraph", behindDoc="0")
    if logo_anchor is not None:
        r_logo._r.getparent().remove(r_logo._r)
        p_main._p.insert(2, logo_anchor)

    r_tl = p_main.add_run()
    r_tl.add_picture(DECO_TL, width=Cm(1.9))
    tl_anchor = make_anchor_from_inline(r_tl, 682625, 1567180, -106680, -106680, 50,
                                        relH="page", relV="paragraph", behindDoc="1")
    if tl_anchor is not None:
        r_tl._r.getparent().remove(r_tl._r)
        p_main._p.insert(3, tl_anchor)

    # Ligne noire au-dessus du nom
    p_l1 = doc.add_paragraph()
    sp(p_l1, before=6, after=4)
    bdr(p_l1, bottom=True, color="231F20", sz="6")

    # Nom + expérience (normal, centré)
    p_nom = doc.add_paragraph()
    p_nom.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp(p_nom, before=4, after=2)
    nom = cv_data.get("nom_prenom", "")
    exp = cv_data.get("annees_experience", "")
    r_n = p_nom.add_run(nom)
    sf(r_n, size_pt=14)
    if exp:
        r_e = p_nom.add_run(f" – {exp}")
        sf(r_e, size_pt=14)

    # Titre poste — gras, centré, taille 14
    p_t = doc.add_paragraph()
    p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp(p_t, before=0, after=4)
    sf(p_t.add_run(cv_data.get("titre_poste", "")), size_pt=14, bold=True)

    # Ligne noire en dessous
    p_l2 = doc.add_paragraph()
    sp(p_l2, before=4, after=8)
    bdr(p_l2, bottom=True, color="231F20", sz="6")


# ─── TITRE DE SECTION ─────────────────────────────────────────────────────────

def section_title(doc, title, page_break=False):
    """
    Titre de section : texte GRAS + ligne noire fine en dessous.
    Comme dans la référence : 'À PROPOS', 'COMPÉTENCES', 'FORMATION'...
    Si page_break=True → saut de page AVANT (pour EXPÉRIENCES PROFESSIONNELLES).
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sp(p, before=10, after=4)
    if page_break:
        page_break_before(p)
    sf(p.add_run(title.upper()), size_pt=10, bold=True, color=BLACK)
    bdr(p, bottom=True, color="231F20", sz="6")
    return p


# ─── À PROPOS ─────────────────────────────────────────────────────────────────

def build_a_propos(doc, a_propos):
    """
    Section À PROPOS — fidèle référence :
    - Chaque bullet = paragraphe séparé avec grand espacement vertical
    - Tiret + espace + texte justifié
    - Mots-clés en gras détectés automatiquement (entre ** ou via liste)
    - Indentation gauche à 426 twips (~0.75cm)
    """
    if not a_propos:
        return

    section_title(doc, "À PROPOS")

    lignes = a_propos if isinstance(a_propos, list) else [a_propos]

    for ligne in lignes:
        ligne = ligne.strip()
        if not ligne:
            continue

        # ── Paragraphe bullet avec espacement référence ──
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        # Espacement important entre bullets comme dans la référence
        sp(p, before=8, after=8, line=276)
        # Indentation : left=426, hanging=0
        ind(p, left=426)

        # Tiret
        r_dash = p.add_run("- ")
        sf(r_dash, size_pt=10, color=BLACK)

        # Traiter les mots en gras (entre ** ou liste de mots-clés)
        _add_text_with_bold(p, ligne)

    # Espace après la section
    p_space = doc.add_paragraph()
    sp(p_space, before=0, after=4)


def _add_text_with_bold(para, text):
    """
    Ajoute du texte dans un paragraphe en gérant le gras.
    Supporte le format **mot** pour le gras, sinon texte simple.
    """
    # Parser les segments **gras** et normal
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            inner = part[2:-2]
            r = para.add_run(inner)
            sf(r, size_pt=10, bold=True, color=BLACK)
        else:
            r = para.add_run(part)
            sf(r, size_pt=10, color=BLACK)


# ─── COMPÉTENCES EN 2 COLONNES ────────────────────────────────────────────────

def build_competences_table(doc, competences):
    """
    Compétences en tableau 2 colonnes SANS bordures visibles.
    - Colonne gauche et droite avec GAP au centre (espace blanc de 0.5cm)
    - Sous-titres en GRAS
    - Items avec tiret et indentation
    - Espacement between catégories
    Fidèle 100% à la référence.
    """
    if not competences:
        return

    section_title(doc, "COMPÉTENCES")

    # Largeurs : page utile = 21 - 3.6 = 17.4cm
    # Col gauche = 8.2cm, gap = 1.0cm, col droite = 8.2cm
    # En twips (1cm = 567 twips)
    col_w   = int(8.2 * 567)   # 4649 twips par colonne
    gap_w   = int(1.0 * 567)   # 567 twips pour le gap
    total_w = col_w * 2 + gap_w

    # Diviser les catégories : moitié gauche, moitié droite
    mid = (len(competences) + 1) // 2
    col_left  = competences[:mid]
    col_right = competences[mid:]

    # Créer tableau 1 ligne x 3 colonnes (col_gauche | gap | col_droite)
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Table Grid'

    # Supprimer les bordures du tableau entier
    tbl_el = tbl._tbl
    tblPr  = tbl_el.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl_el.insert(0, tblPr)

    # Largeur totale du tableau
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'),    str(total_w))
    tblW.set(qn('w:type'), 'dxa')
    old_tblW = tblPr.find(qn('w:tblW'))
    if old_tblW is not None:
        tblPr.remove(old_tblW)
    tblPr.append(tblW)

    # Supprimer bordures tableau
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'),   'none')
        b.set(qn('w:sz'),    '0')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), 'auto')
        tblBorders.append(b)
    old_bdr = tblPr.find(qn('w:tblBorders'))
    if old_bdr is not None:
        tblPr.remove(old_bdr)
    tblPr.append(tblBorders)

    cells = tbl.rows[0].cells

    # Définir largeurs des 3 cellules
    set_cell_width(cells[0], col_w)   # colonne gauche
    set_cell_width(cells[1], gap_w)   # gap vide
    set_cell_width(cells[2], col_w)   # colonne droite

    # Supprimer bordures de chaque cellule
    for cell in cells:
        remove_cell_borders(cell)

    # Remplir les colonnes
    _fill_competence_column(cells[0], col_left)
    _fill_competence_column(cells[2], col_right)
    # Cellule du milieu reste vide

    # Espace après tableau
    p_after = doc.add_paragraph()
    sp(p_after, before=4, after=4)


def _fill_competence_column(cell, categories):
    """Remplit une cellule avec les catégories de compétences."""
    # Vider le paragraphe vide par défaut
    for p in cell.paragraphs:
        p.clear()

    first = True
    for cat in categories:
        # ── Sous-titre catégorie ──
        if cat.get("categorie"):
            if first:
                p = cell.paragraphs[0]
                first = False
            else:
                p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            # Espacement before plus grand entre catégories sauf la première
            sp(p, before=6 if not first else 0, after=2)
            sf(p.add_run(cat["categorie"]), size_pt=10, bold=True, color=BLACK)

        # ── Items avec tiret ──
        for item in cat.get("items", []):
            if first:
                p = cell.paragraphs[0]
                first = False
            else:
                p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            sp(p, before=0, after=0, line=240)
            ind(p, left=283)  # ~0.5cm indentation comme dans la référence
            r_dash = p.add_run("- ")
            sf(r_dash, size_pt=10, color=BLACK)
            r_item = p.add_run(item)
            sf(r_item, size_pt=10, color=BLACK)


# ─── EXPÉRIENCES ──────────────────────────────────────────────────────────────

def exp_section_title(doc):
    """
    Titre EXPÉRIENCES PROFESSIONNELLES :
    - NOUVELLE PAGE avant (page break)
    - Centré, rouge/bordeaux, entre deux lignes noires
    Fidèle 100% à la référence.
    """
    # Ligne noire au-dessus
    p_top = doc.add_paragraph()
    sp(p_top, before=10, after=0)
    # PAGE BREAK BEFORE sur ce paragraphe
    page_break_before(p_top)
    bdr(p_top, bottom=True, color="231F20", sz="6")

    # Titre centré rouge
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp(p, before=4, after=4)
    sf(p.add_run("EXPÉRIENCES PROFESSIONNELLES"), size_pt=12, bold=True, color=TITLE)

    # Ligne noire en dessous
    p_bot = doc.add_paragraph()
    sp(p_bot, before=0, after=8)
    bdr(p_bot, bottom=True, color="231F20", sz="6")


def exp_badge_line(doc, periode, entreprise, badge_id=100):
    """
    Ligne badge + entreprise :
    - Badge rouge arrondi à gauche avec la période en blanc
    - Entreprise en rouge GRAS MAJUSCULES à droite
    - Indentation 1800 twips pour laisser la place au badge

    IMPORTANT : le <w:drawing> du badge DOIT être dans un <w:r> (run),
    pas directement dans <w:p>, sinon Word l'ignore silencieusement.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sp(p, before=10, after=2)
    ind(p, left=1800)

    if periode:
        try:
            badge = make_badge_xml(periode, doc_id=badge_id)
            # ── Encapsuler le badge dans un <w:r> ──────────────────────────
            r_badge = OxmlElement('w:r')
            r_badge.append(badge)
            # Insérer après <w:pPr> (index 1)
            pPr = p._p.find(qn('w:pPr'))
            if pPr is not None:
                pPr.addnext(r_badge)
            else:
                p._p.insert(0, r_badge)
        except Exception as e:
            # Fallback texte si badge échoue
            r_per = p.add_run(f"[{periode}]  ")
            sf(r_per, size_pt=9, bold=True, color=RED)

    if entreprise:
        sf(p.add_run(entreprise.upper()), size_pt=10, bold=True, color=RED)

    return p


def exp_poste(doc, poste):
    """Poste centré gras sous le badge — comme dans la référence."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp(p, before=2, after=6)
    sf(p.add_run(poste), size_pt=10, bold=True, color=BLACK)
    return p


def exp_label(doc, label):
    """
    Label de sous-section : GRAS, aligné gauche.
    Ex: 'Contexte & enjeux :', 'Objectifs :', 'Réalisations :', 'Environnement :'
    PAS de ligne de séparation — seulement espacement.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sp(p, before=8, after=2)
    sf(p.add_run(label), size_pt=10, bold=True, color=BLACK)
    return p


def exp_body(doc, text, indent_twips=426):
    """Corps de texte justifié avec tiret, pour Contexte et Environnement."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    sp(p, before=0, after=2, line=252)
    ind(p, left=indent_twips)
    r_dash = p.add_run("- ")
    sf(r_dash, size_pt=10, color=BLACK)
    sf(p.add_run(text), size_pt=10, color=BLACK)
    return p


def exp_item(doc, text, indent_twips=426):
    """
    Item de liste simple avec tiret.
    Pour: Objectifs, items sans gras initial.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    sp(p, before=1, after=1, line=252)
    ind(p, left=indent_twips)
    r_dash = p.add_run("- ")
    sf(r_dash, size_pt=10, color=BLACK)
    sf(p.add_run(text), size_pt=10, color=BLACK)
    return p


def exp_item_bold_prefix(doc, text, indent_twips=426):
    """
    Item de réalisation avec PREMIER TERME EN GRAS jusqu'au ':'.
    Ex: '- Optimisation & Diagnostic JVM : Analyse...'
        → '- ' + 'Optimisation & Diagnostic JVM :' (gras) + ' Analyse...' (normal)
    Fidèle 100% à la référence.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    sp(p, before=1, after=1, line=252)
    ind(p, left=indent_twips)

    # Tiret
    r_dash = p.add_run("- ")
    sf(r_dash, size_pt=10, color=BLACK)

    # Chercher le séparateur ':' ou '–'
    colon_idx = text.find(' : ')
    dash_idx  = text.find(' – ')

    if colon_idx > 0:
        bold_part  = text[:colon_idx + 2]   # inclut ' :'
        normal_part = text[colon_idx + 2:]   # reste après ' :'
    elif dash_idx > 0:
        bold_part  = text[:dash_idx + 2]
        normal_part = text[dash_idx + 2:]
    else:
        # Pas de séparateur → tout normal
        sf(p.add_run(text), size_pt=10, color=BLACK)
        return p

    r_bold = p.add_run(bold_part)
    sf(r_bold, size_pt=10, bold=True, color=BLACK)

    if normal_part.strip():
        r_normal = p.add_run(normal_part)
        sf(r_normal, size_pt=10, color=BLACK)

    return p


# ─── PIED DE PAGE ─────────────────────────────────────────────────────────────

def build_footer(doc):
    section = doc.sections[0]
    footer  = section.footer
    for para in footer.paragraphs:
        para.clear()
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sp(fp, before=0, after=0)

    pPr  = fp._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top  = OxmlElement('w:top')
    top.set(qn('w:val'),   'single')
    top.set(qn('w:sz'),    '6')
    top.set(qn('w:space'), '1')
    top.set(qn('w:color'), 'auto')
    pBdr.append(top)
    pPr.append(pBdr)

    tabs_el = OxmlElement('w:tabs')
    for val, pos in [('center', '4536'), ('right', '9026')]:
        t = OxmlElement('w:tab')
        t.set(qn('w:val'), val)
        t.set(qn('w:pos'), pos)
        tabs_el.append(t)
    pPr.append(tabs_el)

    sf(fp.add_run("MALTEM AFRICA"), size_pt=8, color=GREY)
    fp.add_run("\t").font.size = Pt(8)

    def field(para, instr):
        r = para.add_run()
        r.font.size = Pt(8)
        r.font.color.rgb = GREY
        f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin')
        it = OxmlElement('w:instrText'); it.text = instr
        f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'end')
        r._r.extend([f1, it, f2])

    sf(fp.add_run("\tPage "), size_pt=8, color=GREY)
    field(fp, ' PAGE ')
    sf(fp.add_run(" sur "), size_pt=8, color=GREY)
    field(fp, ' NUMPAGES ')

    # Deco bas-droite
    r_br = fp.add_run()
    r_br.add_picture(DECO_BR, width=Cm(1.87))
    br_anchor = make_anchor_from_inline(r_br, 673100, 1567180, 6886900, -1200000, 51,
                                        relH="page", relV="paragraph", behindDoc="1")
    if br_anchor is not None:
        r_br._r.getparent().remove(r_br._r)
        fp._p.insert(1, br_anchor)

    # Deco haut-gauche (depuis footer, remonte en haut de page)
    r_tl = fp.add_run()
    r_tl.add_picture(DECO_TL, width=Cm(1.9))
    tl_anchor = make_anchor_from_inline(r_tl, 682625, 1567180, -106680, -9500000, 52,
                                        relH="page", relV="paragraph", behindDoc="1")
    if tl_anchor is not None:
        r_tl._r.getparent().remove(r_tl._r)
        fp._p.insert(1, tl_anchor)


# ─── GÉNÉRATION PRINCIPALE ────────────────────────────────────────────────────

def generate_maltem_cv(cv_data: dict, output_path: str) -> str:
    doc = Document()
    s   = doc.sections[0]
    s.page_height      = Cm(29.7)
    s.page_width       = Cm(21.0)
    s.left_margin      = Cm(1.8)
    s.right_margin     = Cm(1.8)
    s.top_margin       = Cm(1.8)
    s.bottom_margin    = Cm(1.8)
    s.footer_distance  = Cm(0.8)
    doc.styles['Normal'].font.name = 'Century Gothic'
    doc.styles['Normal'].font.size = Pt(10)

    # ── EN-TÊTE ──────────────────────────────────────────────────────────────
    build_header(doc, cv_data)

    # ── À PROPOS ─────────────────────────────────────────────────────────────
    a_propos = cv_data.get("a_propos", "")
    if a_propos:
        build_a_propos(doc, a_propos)

    # ── COMPÉTENCES (2 colonnes avec gap) ────────────────────────────────────
    competences = cv_data.get("competences", [])
    if competences:
        build_competences_table(doc, competences)

    # ── FORMATION ────────────────────────────────────────────────────────────
    formations = cv_data.get("formations", [])
    if formations:
        section_title(doc, "FORMATION")
        for form in formations:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            sp(p, before=4, after=2)
            annee       = form.get("annee", "")
            diplome     = form.get("diplome", "")
            etablissement = form.get("etablissement", "")
            if annee:
                r_yr = p.add_run(f"{annee}  ")
                sf(r_yr, size_pt=10, bold=True, color=BLACK)
            txt = diplome
            if etablissement:
                txt += f" – {etablissement}"
            sf(p.add_run(txt), size_pt=10, color=BLACK)

    # ── CERTIFICATIONS ───────────────────────────────────────────────────────
    certs = cv_data.get("certifications", [])
    if certs:
        section_title(doc, "CERTIFICATIONS")
        for cert in certs:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            sp(p, before=0, after=0)
            ind(p, left=426)
            r_dash = p.add_run("- ")
            sf(r_dash, size_pt=10, color=BLACK)
            sf(p.add_run(cert), size_pt=10, color=BLACK)

    # ── LANGUES ──────────────────────────────────────────────────────────────
    langues = cv_data.get("langues", [])
    if langues:
        section_title(doc, "LANGUES")
        for langue in langues:
            p = doc.add_paragraph()
            sp(p, before=0, after=0)
            ind(p, left=426)
            r_dash = p.add_run("- ")
            sf(r_dash, size_pt=10, color=BLACK)
            sf(p.add_run(langue), size_pt=10, color=BLACK)

    # ── AUTRES RÉFÉRENCES ────────────────────────────────────────────────────
    refs = cv_data.get("autres_references", [])
    if refs:
        section_title(doc, "AUTRES RÉFÉRENCES")
        for ref in refs:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            sp(p, before=2, after=2)
            # Entreprise en rouge gras + poste en noir
            entreprise_ref = ref.get("entreprise", "") if isinstance(ref, dict) else ""
            poste_ref      = ref.get("poste", "")      if isinstance(ref, dict) else ""
            if entreprise_ref:
                r_ent = p.add_run(f"{entreprise_ref} : ")
                sf(r_ent, size_pt=10, bold=True, color=RED)
            if poste_ref:
                sf(p.add_run(poste_ref), size_pt=10, color=BLACK)
            # Si c'est juste une string
            if isinstance(ref, str):
                sf(p.add_run(ref), size_pt=10, color=BLACK)

    # ── PROJETS MARQUANTS ────────────────────────────────────────────────────
    projets = cv_data.get("projets_marquants", [])
    if projets:
        section_title(doc, "PROJETS MARQUANTS")
        for projet in projets:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            sp(p, before=1, after=1)
            ind(p, left=426)
            r_dash = p.add_run("- ")
            sf(r_dash, size_pt=10, color=BLACK)
            sf(p.add_run(projet if isinstance(projet, str) else str(projet)), size_pt=10, color=BLACK)

    # ── EXPÉRIENCES PROFESSIONNELLES (nouvelle page) ──────────────────────────
    experiences = cv_data.get("experiences", [])
    if experiences:
        # Titre avec saut de page avant
        exp_section_title(doc)

        badge_id = 100
        for exp in experiences:
            periode       = exp.get("periode", "")
            entreprise    = exp.get("entreprise", "")
            poste         = exp.get("poste", "")
            direction     = exp.get("direction", "")
            contexte      = exp.get("contexte", "")
            objectifs     = exp.get("objectifs", [])
            missions      = exp.get("missions", [])
            realisations  = exp.get("realisations", [])
            resultats     = exp.get("resultats", [])
            environnement = exp.get("environnement", "")

            # Badge période + entreprise
            exp_badge_line(doc, periode, entreprise, badge_id=badge_id)
            badge_id += 1

            # Poste centré gras
            if poste:
                exp_poste(doc, poste)

            # Direction (si présente)
            if direction:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                sp(p, before=0, after=2)
                sf(p.add_run(direction), size_pt=10, bold=True, color=BLACK)

            # Contexte & enjeux
            if contexte:
                exp_label(doc, "Contexte & enjeux :")
                if isinstance(contexte, list):
                    for c in contexte:
                        exp_item(doc, c)
                else:
                    exp_body(doc, contexte)

            # Objectifs
            if objectifs:
                exp_label(doc, "Objectifs :")
                for o in objectifs:
                    exp_item(doc, o)

            # Réalisations (premier terme en gras)
            reals = realisations if realisations else missions
            if reals:
                exp_label(doc, "Réalisations :")
                for r in reals:
                    exp_item_bold_prefix(doc, r)

            # Résultats / impacts
            if resultats:
                exp_label(doc, "Résultats / impacts :")
                for r in resultats:
                    exp_item(doc, r)

            # Environnement
            if environnement:
                exp_label(doc, "Environnement :")
                if isinstance(environnement, list):
                    for e in environnement:
                        exp_item(doc, e)
                else:
                    exp_body(doc, environnement)

            # Séparateur entre expériences
            p_sep = doc.add_paragraph()
            sp(p_sep, before=6, after=0)

    # ── PIED DE PAGE ─────────────────────────────────────────────────────────
    build_footer(doc)

    doc.save(output_path)
    return output_path
