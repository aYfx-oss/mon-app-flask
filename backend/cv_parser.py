"""
cv_parser.py — Extraction du texte brut depuis un CV PDF ou DOCX
"""
import os
import pdfplumber
from docx import Document


def extract_text_from_pdf(filepath: str) -> str:
    """Extrait le texte d'un fichier PDF."""
    text = []
    with pdfplumber.open(filepath) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text.append(page_text)
    return "\n".join(text)


def extract_text_from_docx(filepath: str) -> str:
    """Extrait le texte d'un fichier DOCX."""
    doc = Document(filepath)
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text.strip())
    # Inclure aussi le texte dans les tableaux
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    return "\n".join(paragraphs)


def extract_cv_text(filepath: str) -> str:
    """Détecte le type de fichier et extrait le texte."""
    ext = os.path.splitext(filepath)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(filepath)
    elif ext in (".docx", ".doc"):
        return extract_text_from_docx(filepath)
    else:
        raise ValueError(f"Format non supporté : {ext}. Utilisez PDF ou DOCX.")
